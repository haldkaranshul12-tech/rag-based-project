"""
Unified text extraction for PDF, DOCX, and image files.

Returns a list of per-page strings (extract_pages) so downstream chunking
can tag every chunk with the page it came from. DOCX and images have no
native page concept, so they are treated as a single page.
"""

from pypdf import PdfReader


def extract_pages_from_pdf(file):
    """Returns a list of strings, one per PDF page (in page order)."""
    reader = PdfReader(file)
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages.append(page_text)

    if sum(len(p.strip()) for p in pages) >= 20:
        return pages

    # No real text layer found on any page (likely a scanned PDF) -- OCR it
    return extract_pages_from_scanned_pdf(file)


def extract_pages_from_scanned_pdf(file):
    try:
        import fitz  # PyMuPDF
        from PIL import Image
        import pytesseract
        import io
    except ImportError:
        raise RuntimeError(
            "OCR support for scanned PDFs needs PyMuPDF, Pillow, and "
            "pytesseract. In your activated venv, run: "
            "pip install PyMuPDF Pillow pytesseract"
        )

    file.seek(0)
    pdf_doc = fitz.open(stream=file.read(), filetype="pdf")

    pages = []
    for page in pdf_doc:
        pix = page.get_pixmap(dpi=200)
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        try:
            pages.append(pytesseract.image_to_string(image))
        except pytesseract.TesseractNotFoundError:
            raise RuntimeError(
                "Tesseract OCR engine not found on this system. Install it "
                "separately (not just via pip) -- see "
                "https://github.com/UB-Mannheim/tesseract/wiki for Windows, "
                "then set pytesseract.pytesseract.tesseract_cmd to its install path."
            )

    return pages


def extract_pages_from_docx(file):
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run this in your activated venv: "
            "pip install python-docx"
        )

    document = Document(file)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for c in row.cells:
                if c.text.strip():
                    paragraphs.append(c.text)

    # DOCX has no reliable page concept without rendering -- treat as one page
    return ["\n".join(paragraphs)]


def extract_pages_from_image(file):
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        raise RuntimeError(
            "pytesseract/Pillow are not installed. Run this in your "
            "activated venv: pip install pytesseract Pillow"
        )

    # If Windows doesn't auto-detect Tesseract, uncomment and set your path:
    # pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

    image = Image.open(file)
    try:
        return [pytesseract.image_to_string(image)]
    except pytesseract.TesseractNotFoundError:
        raise RuntimeError(
            "Tesseract OCR engine not found on this system. Install it "
            "separately (not just via pip) -- see "
            "https://github.com/UB-Mannheim/tesseract/wiki for Windows, "
            "then set pytesseract.pytesseract.tesseract_cmd to its install path."
        )


def get_doc_type(filename):
    """Returns a short document-type label based on the file extension."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return "pdf"
    elif name.endswith(".docx"):
        return "docx"
    elif name.endswith((".png", ".jpg", ".jpeg")):
        return "image"
    return "unknown"


def extract_pages(uploaded_file):
    """
    Looks at the uploaded file's extension and returns a list of per-page
    text strings. `uploaded_file` is the Streamlit UploadedFile object.
    """
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return extract_pages_from_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return extract_pages_from_docx(uploaded_file)
    elif name.endswith((".png", ".jpg", ".jpeg")):
        return extract_pages_from_image(uploaded_file)
    else:
        raise ValueError(f"Unsupported file type: {uploaded_file.name}")


def extract_text(uploaded_file):
    """Backwards-compatible helper: full document text as one string."""
    return "\n".join(extract_pages(uploaded_file))